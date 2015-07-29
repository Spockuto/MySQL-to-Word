#!/bin/bash
import MySQLdb
import os
import getpass
from docx import Document

username=raw_input("User:")
password=getpass.getpass()
dbname=raw_input("Database:") 
tablename=raw_input("Table:")

db=MySQLdb.connect(host='localhost',user='%s'%username,passwd='%s'%password,db='%s'%dbname)
cursor=db.cursor()

sql='''SELECT COUNT(*) from %s'''%tablename
cursor.execute(sql)
(rows,)=cursor.fetchone()
rows=int(rows)+1

sql='''SELECT * from %s'''%tablename
cursor.execute(sql)
data=cursor.fetchall()
column=len(cursor.description)

document=Document()
document.add_heading('%s'%tablename, 0)
table = document.add_table(rows=int(rows),cols=int(column))
header = table.rows[0].cells

for cols in range(0,column):
	header[cols].text = str(cursor.description[cols][0])

for row in range(1,rows):
	data_rows = table.rows[row].cells
	for cols in range(0,column):
		data_rows[cols].text = str(data[row-1][cols])

db.close()
document.save('%s-%s.docx'%(dbname,tablename))
